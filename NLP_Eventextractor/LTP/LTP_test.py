import os
import docx
import re
from docx import Document
from ltp import LTP
ltp = LTP(path='base')


path = r'D:\pycharm\untitled3\NLP\ProvincePolice_test'#文档地址

user_dict_path = r'D:\pycharm\untitled3\NLP\User_dict\user_dict.txt'  #词典地址
#自定义词典
# user_dict.txt 是词典文件， max_window是最大前向分词窗口
ltp.init_dict(path=user_dict_path, max_window=4)
# 也可以在代码中添加自定义的词语
ltp.add_words(words=["博时基金", "长江大桥"], max_window=4)



#分句

sents = ltp.sent_split(["他叫汤姆去拿外衣。", "汤姆生病了。他去了医院。"])
print(sents)

#词性标注
seg, hidden = ltp.seg(["博时基金建议行业配置关注三条主线"])
pos = ltp.pos(hidden)
print(seg)
print(pos)


'''
#命名实体识别
ner = ltp.ner(hidden)
# [['他', '叫', '汤姆', '去', '拿', '外衣', '。']]
# [[('Nh', 2, 2)]]
print(ner)
tag, start, end = ner[0][0]
print(ner)
print(tag,":", "".join(seg[0][start:end + 1]))
# Nh : 汤姆

'''

#依存句法分析
srl = ltp.srl(hidden)
# [['他', '叫', '汤姆', '去', '拿', '外衣', '。']]
# [
#     [
#         [],                                                # 他
#         [('ARG0', 0, 0), ('ARG1', 2, 2), ('ARG2', 3, 5)],  # 叫 -> [ARG0: 他, ARG1: 汤姆, ARG2: 去拿外衣]
#         [],                                                # 汤姆
#         [],                                                # 去
#         [('ARG0', 2, 2), ('ARG1', 5, 5)],                  # 拿 -> [ARG0: 汤姆, ARG1: 外衣]
#         [],                                                # 外衣
#         []                                                 # 。
#     ]
# ]
srl = ltp.srl(hidden, keep_empty=False)
# [
#     [
#         (1, [('ARG0', 0, 0), ('ARG1', 2, 2), ('ARG2', 3, 5)]), # 叫 -> [ARG0: 他, ARG1: 汤姆, ARG2: 去拿外衣]
#         (4, [('ARG0', 2, 2), ('ARG1', 5, 5)])                  # 拿 -> [ARG0: 汤姆, ARG1: 外衣]
#     ]
# ]
print(srl)

#语义依存分析(树)
sdp_tree = ltp.sdp(hidden, mode='tree')
# [['他', '叫', '汤姆', '去', '拿', '外衣', '。']]
# [
#     [
#         (1, 2, 'Agt'),
#         (2, 0, 'Root'),   # 叫 --|Root|--> ROOT
#         (3, 2, 'Datv'),
#         (4, 2, 'eEfft'),
#         (5, 4, 'eEfft'),
#         (6, 5, 'Pat'),
#         (7, 2, 'mPunc')
#     ]
# ]
print(sdp_tree)

#语义依存分析(图)
sdp_graph = ltp.sdp(hidden, mode='graph')
# [['他', '叫', '汤姆', '去', '拿', '外衣', '。']]
# [
#     [
#         (1, 2, 'Agt'),
#         (2, 0, 'Root'),   # 叫 --|Root|--> ROOT
#         (3, 2, 'Datv'),
#         (3, 4, 'Agt'),
#         (3, 5, 'Agt'),
#         (4, 2, 'eEfft'),
#         (5, 4, 'eEfft'),
#         (6, 5, 'Pat'),
#         (7, 2, 'mPunc')
#     ]
# ]
print(sdp_graph)


'''
if __name__ == '__main__':

    l = os.listdir(path)  # 把文件夹中的文件名以文本的形式放入列表
    cout = 0
    for i in l:  # 遍历每一个文件名
        c = ''
        if 'docx' not in i:  # 跳过不是docx的文件
            continue
        fi_d = os.path.join(path, i)
        cout += 1
        print("第（%d）份文件位于" % cout, fi_d)  # 打印文件位置及名称
        document = Document(str(fi_d))
        # print(document.paragraphs[0].text)
        key = re.compile(r"完善农业.*", re.S)  # 匹配相应内容
        m = 0

        for paragraph in document.paragraphs:

            m = m + 1
            te = re.findall(key, paragraph.text)  # 正则匹配符合条件的段落
            strte = ''.join(te)  # 将返回的列表转换成字符串
            if strte == '':  # 若未匹配到符合条件的段落，则跳出本次循环
                continue
            
            c = c + paragraph.text
            strte = c
           
            # 词性标注
            seg, hidden = ltp.seg([strte])
            pos = ltp.pos(hidden)
            
            s_sum = sum(seg,[])
            p_sum = sum(pos,[])

            s_p = dict(zip(s_sum,p_sum))
            #print(p)
            for s,p in s_p.items():
                print(s,p)
            
            #print(pos)
            #print(strte)
            # 命名实体识别
            ner = ltp.ner(hidden)
            # [['他', '叫', '汤姆', '去', '拿', '外衣', '。']]
            # [[('Nh', 2, 2)]]

            tag, start, end = ner[0][0]
            print(tag, ":", "".join(seg[0][start:end + 1]))
            # Nh : 汤姆

print('\n\n')
print("共有 %d 份文件"%cout)
'''